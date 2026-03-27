"""
SPECTRA Documentation Generator
Produces a fully formatted Word document covering:
  1. SPECTRA Overview & Architecture
  2. Page-by-Page Computation Reference
  3. Role-Based Usage Guide (Risk Officer / Admin / Analyst)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── helpers ───────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F3F4F6')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 64, 175)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1E3A5F')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        fill = 'F8FAFC' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), fill)
            cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CBD5E1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_callout(doc, text, bg='EFF6FF', border='BFDBFE', label=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), bg)
    p._p.get_or_add_pPr().append(shading)
    if label:
        r = p.add_run(label + '  ')
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(30, 64, 175)
    p.add_run(text).font.size = Pt(10)
    return p

# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Heading styles
for lvl, sz, color in [
    (1, 20, (15, 23, 42)),
    (2, 15, (30, 58, 138)),
    (3, 13, (71, 85, 105)),
    (4, 11, (71, 85, 105)),
]:
    hstyle = doc.styles[f'Heading {lvl}']
    hstyle.font.name = 'Calibri'
    hstyle.font.size = Pt(sz)
    hstyle.font.bold = True
    hstyle.font.color.rgb = RGBColor(*color)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SPECTRA')
run.font.name = 'Calibri'
run.font.size = Pt(48)
run.font.bold = True
run.font.color.rgb = RGBColor(15, 23, 42)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('Credit Risk Intelligence Platform')
run2.font.size = Pt(20)
run2.font.color.rgb = RGBColor(71, 85, 105)

doc.add_paragraph()
line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = line.add_run('─' * 60)
run3.font.color.rgb = RGBColor(148, 163, 184)

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Complete System Reference  ·  Architecture  ·  Computation Logic  ·  Role Guides')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(f'Generated: {datetime.datetime.now().strftime("%d %B %Y")}').font.color.rgb = RGBColor(148, 163, 184)

ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver_p.add_run('Version 1.0  ·  Confidential — Internal Use Only').font.color.rgb = RGBColor(148, 163, 184)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Table of Contents', 1)
toc_entries = [
    ('PART I — SPECTRA OVERVIEW & ARCHITECTURE', ''),
    ('  1.  What SPECTRA Is', ''),
    ('  2.  The Problem It Solves', ''),
    ('  3.  Core Features', ''),
    ('  4.  How It Works — Data Flow', ''),
    ('  5.  Technical Architecture', ''),
    ('  6.  Database Tables', ''),
    ('  7.  AI & Analytics — Predictions & Models', ''),
    ('  8.  API Endpoints Reference', ''),
    ('  9.  Use Cases', ''),
    ('  10. Outputs & Insights', ''),
    ('  11. Benefits', ''),
    ('  12. Example Scenario', ''),
    ('', ''),
    ('PART II — PAGE-BY-PAGE COMPUTATION REFERENCE', ''),
    ('  13. Foundation Layer (DB, Cache, Config)', ''),
    ('  14. Dashboard Page', ''),
    ('  15. Portfolio Page', ''),
    ('  16. Clients List Page', ''),
    ('  17. Client Profile Page', ''),
    ('  18. Analytics Page', ''),
    ('  19. Watchlist Page', ''),
    ('  20. Stress Test Page', ''),
    ('  21. Concentration Risk Page', ''),
    ('  22. Model Intelligence Page', ''),
    ('  23. Audit Log Page', ''),
    ('  24. The Classification Engine', ''),
    ('  25. The Action Engine', ''),
    ('  26. The Notification Service', ''),
    ('', ''),
    ('PART III — ROLE-BASED USAGE GUIDES', ''),
    ('  27. Risk Officer Guide', ''),
    ('  28. Admin Guide', ''),
    ('  29. Analyst Guide', ''),
    ('  30. Role Comparison Table', ''),
]
for entry, _ in toc_entries:
    p = doc.add_paragraph()
    if entry == '':
        continue
    if not entry.startswith(' '):
        run = p.add_run(entry)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(15, 23, 42)
    else:
        run = p.add_run(entry)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(71, 85, 105)
        p.paragraph_format.left_indent = Inches(0.2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PART I — SPECTRA OVERVIEW & ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'PART I — SPECTRA OVERVIEW & ARCHITECTURE', 1)
add_divider(doc)

# ── 1. What SPECTRA Is ───────────────────────────────────────────────────────
add_heading(doc, '1.  What SPECTRA Is', 2)

add_para(doc,
    'SPECTRA (Systematic Portfolio Early Credit Threat Risk Analytics) is a credit risk intelligence '
    'platform built for banks, fintechs, and microfinance institutions. It sits on top of an existing '
    'core banking SQL Server database and transforms raw transactional data into actionable, real-time '
    'risk decisions — without modifying any origination data.')

add_para(doc,
    'SPECTRA combines four disciplines into a single workspace:')
add_bullet(doc, 'IFRS 9 — automatic, continuous impairment staging (Stage 1 / 2 / 3)')
add_bullet(doc, 'Machine Learning — probability of default (PD) predictions at 30, 60, and 90-day horizons')
add_bullet(doc, 'Early Warning Indicators (EWI) — real-time behavioural signals (salary, overdraft, card, DPD)')
add_bullet(doc, 'Case Management — engagement logging, document requests, covenant waivers, restructuring plans')

# ── 2. The Problem It Solves ─────────────────────────────────────────────────
add_heading(doc, '2.  The Problem It Solves', 2)

add_para(doc, 'Traditional banks manage credit risk reactively:')
add_bullet(doc, 'Stage classification is done in monthly batch runs — problems are discovered too late')
add_bullet(doc, 'Relationship Managers (RMs) work from spreadsheets and email — no single source of truth')
add_bullet(doc, 'Credit committee decisions lack quantitative backing — relying on qualitative judgment alone')
add_bullet(doc, 'There is no systematic escalation path — deteriorating accounts fall through the cracks')
add_bullet(doc, 'Audit trails are fragmented across systems — regulatory compliance is difficult to demonstrate')

doc.add_paragraph()
add_para(doc, 'SPECTRA solves all of these:')
add_bullet(doc, 'Every client is scored and staged daily — not monthly')
add_bullet(doc, 'Alerts fire automatically the moment a threshold is crossed')
add_bullet(doc, 'Every RM action is logged, timestamped, attributed, and auditable')
add_bullet(doc, 'AI generates risk narratives and transparency letters on demand')
add_bullet(doc, 'Credit committee receives quantitative evidence (PD scores, SHAP drivers, stress scenarios)')

# ── 3. Core Features ─────────────────────────────────────────────────────────
add_heading(doc, '3.  Core Features', 2)

features = [
    ('IFRS 9 Risk Classification',
     'Automatic assignment to Stage 1 (performing), Stage 2 (SICR — Significant Increase in Credit Risk), '
     'or Stage 3 (NPL / default). Stage changes trigger notifications, credit freezes, and review frequency '
     'adjustments automatically. Stage never auto-downgrades.'),
    ('Machine Learning Predictions',
     'Probability of Default (PD) scores at 30, 60, and 90-day horizons. SHAP explanations show the top 3 '
     'factors driving each client\'s score. Six binary risk flags detect behavioural anomalies.'),
    ('Early Warning Indicators (EWI)',
     'Monitors salary inflow, overdraft usage, card utilisation, and consecutive late payments in real time. '
     'EWI signals can be fired via API, triggering immediate reclassification and notifications.'),
    ('Client Engagement Management',
     'RMs log calls and meetings with outcome tracking (reached / no answer / productive / inconclusive). '
     'Full history per client with timestamps and attribution.'),
    ('Document & Collateral Tracking',
     'Request financial statements, bank statements, tax returns — tracked from Pending to Received. '
     'Collateral revaluations with LTV auto-calculation server-side.'),
    ('Covenant Waivers',
     'Formal waiver lifecycle: Pending → Approved / Rejected. Gated to risk_officer and admin roles. '
     'Decision notes recorded permanently.'),
    ('Restructuring Plans',
     'Full proposal-to-completion lifecycle across five plan types: Loan Extension, Payment Holiday, '
     'Rate Reduction, Debt Consolidation, Partial Write-Off.'),
    ('AI-Generated Insights',
     'On-demand risk narrative, deterioration prediction, recommended actions, recovery strategy, and '
     'client transparency letter — generated by Claude AI (claude-haiku-4-5).'),
    ('Notifications & Audit Trail',
     'RM inbox with unread count. Every stage change and EWI trigger creates a notification. '
     'Immutable SystemActions log plus manual ClientActions log.'),
    ('Portfolio Analytics',
     'NPL ratio, Stage 2 rate, concentration risk (HHI), vintage analysis, ECL provision gap, '
     'stress testing with Baseline / Adverse / Severe scenarios.'),
]

for name, desc in features:
    add_bullet(doc, f' — {desc}', bold_prefix=name)

# ── 4. How It Works — Data Flow ──────────────────────────────────────────────
add_heading(doc, '4.  How It Works — Data Flow', 2)

add_para(doc, 'Data flows through SPECTRA in four layers:')
add_numbered(doc, 'Core Banking DB (SQL Server) — bank-originated, read-only tables feed daily snapshots',
             bold_prefix='Source Layer: ')
add_numbered(doc, 'Python ML Pipeline (nightly) — produces PD scores, SHAP explanations, and risk flags as CSV files',
             bold_prefix='ML Layer: ')
add_numbered(doc, 'SPECTRA API (Next.js/Node.js) — classification engine, action engine, and REST endpoints',
             bold_prefix='Logic Layer: ')
add_numbered(doc, 'React Frontend — server-rendered pages + client-side interactive state',
             bold_prefix='Presentation Layer: ')

doc.add_paragraph()
add_callout(doc,
    'The bank\'s source tables (Customer, Credits, RiskPortfolio, DueDaysDaily, etc.) are READ-ONLY. '
    'SPECTRA never modifies origination data. All SPECTRA-owned state lives in separately managed tables '
    'that are created automatically on first API call (DDL-on-first-use pattern).',
    bg='FEF3C7', label='⚠  Key Design Principle:')

# ── 5. Technical Architecture ────────────────────────────────────────────────
add_heading(doc, '5.  Technical Architecture', 2)

add_heading(doc, 'Stack', 3)
add_table(doc,
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Frontend', 'Next.js 14 + React + TypeScript', 'Server-rendered pages, client-side tabs and modals'],
        ['Backend', 'Next.js API Routes (Node.js)', '~30 REST endpoints, JWT auth, role-based access control'],
        ['Database', 'SQL Server 2016+', 'Source data (read-only) + SPECTRA-managed tables (DDL-on-first-use)'],
        ['ML Pipeline', 'Python (external, nightly)', 'PD scoring, SHAP, risk flags — outputs to CSV'],
        ['AI', 'Anthropic Claude API (claude-haiku-4-5)', 'On-demand risk narratives and transparency letters'],
        ['Auth', 'JWT HS256, 8-hour sessions', 'Secure httpOnly cookies, rate-limited login (5 attempts)'],
        ['Caching', 'Module-level in-process cache (LRU, 200 entries)', '5-min TTL for DB queries, 15-min for CSV files'],
    ],
    col_widths=[1.2, 2.2, 2.8]
)

add_heading(doc, 'Connection Pool', 3)
add_bullet(doc, 'min: 2 warm connections always open — first request never waits for TCP handshake')
add_bullet(doc, 'max: 20 connections — handles burst concurrency for multi-user environment')
add_bullet(doc, 'idleTimeout: 60 seconds — keeps connections warm between page navigations')
add_bullet(doc, 'acquireTimeout: 15 seconds — fails fast if all 20 connections are busy')
add_bullet(doc, 'Per-query timeout: 30 seconds default — set per request, overrides ODBC-level timeout')

add_heading(doc, 'DDL-on-First-Use Pattern', 3)
add_para(doc,
    'Every SPECTRA-owned table is created automatically on the first API call that needs it. '
    'A module-level guard (_tablesReady / _tablesInFlight) ensures: '
    '(1) the DDL runs exactly once even under concurrent load, '
    '(2) every subsequent call skips the DDL entirely via a fast boolean check, '
    '(3) no migration scripts or deployment steps are required.')

# ── 6. Database Tables ───────────────────────────────────────────────────────
add_heading(doc, '6.  Database Tables', 2)

add_heading(doc, '6.1  Source Tables (Read-Only — Bank Originated)', 3)
add_table(doc,
    ['Table', 'Key Columns', 'Purpose'],
    [
        ['Customer', 'PersonalID, name, surname, DOB, Gender, City, Branch, Occupation', 'Client personal data'],
        ['Credits', 'CreditAccount, NoCredit, TypeOfCalculatioin, Amount, Interes, Period, FromYear, ToYear, STATUS', 'Loan product details'],
        ['RiskPortfolio', 'clientID, contractNumber, Stage, totalExposure, CalculationDate, TypeOfProduct', 'Daily IFRS 9 staging and exposure'],
        ['DueDaysDaily', 'PersonalID, CreditAccount, DueDays, dateID', 'Daily DPD snapshots'],
        ['AmortizationPlan', 'PARTIJA, RB, DATUMDOSPECA, IZNOS, KAMATA, ANUITET, OTPLATA', 'Scheduled payment plan'],
        ['Accounts / TAccounts', 'PersonalID, Account, Balance / Amount, Date', 'Transaction accounts and monthly transactions'],
        ['Cards / CC_Event_LOG', 'PersonalID, NoCards, brand, type / Account, Amount, trans_date', 'Card products and spend events'],
    ],
    col_widths=[1.5, 2.5, 2.2]
)

add_heading(doc, '6.2  SPECTRA-Managed Tables (DDL-on-First-Use)', 3)
add_table(doc,
    ['Table', 'Key Columns', 'Purpose'],
    [
        ['ClientActions', 'id, clientId, action, status, actionedBy, notes, createdAt', 'Manual RM action log'],
        ['SystemActions', 'id, client_id, event_type, old_stage, new_stage, trigger_reason, created_at', 'Immutable automated event audit log'],
        ['Notifications', 'id, client_id, notification_type, priority, title, message, assigned_rm, read_at', 'RM inbox alerts'],
        ['ClientMonitoring', 'client_id PK, review_frequency, is_freezed, freeze_reason, frozen_at', 'Auto-updated on stage change'],
        ['DocumentRequests', 'id, client_id, document_type, status (Pending/Received), requested_by, received_at', 'Document collection tracking'],
        ['CollateralReview', 'id, client_id, revaluation_date, old_value, new_value, current_exposure, ltv_recalculated', 'Collateral revaluation history'],
        ['ClientEngagements', 'id, client_id, type (call/meeting), scheduled_at, status, outcome, logged_by', 'RM call and meeting log'],
        ['CovenantWaivers', 'id, client_id, waiver_type, status (Pending/Approved/Rejected), approved_by, decision_notes', 'Formal waiver requests'],
        ['RestructuringPlans', 'id, client_id, type, new_maturity_date, holiday_duration_months, new_interest_rate, forgiven_amount, status', 'Loan restructuring proposals'],
    ],
    col_widths=[1.8, 2.8, 1.6]
)

# ── 7. AI & Analytics ────────────────────────────────────────────────────────
add_heading(doc, '7.  AI & Analytics — How Predictions Are Made', 2)

add_heading(doc, '7.1  Machine Learning Pipeline (Python, Nightly)', 3)
add_para(doc, 'The Python pipeline runs nightly and produces three CSV files:')
add_table(doc,
    ['File', 'Content', 'Cache TTL'],
    [
        ['predictions.csv', 'pd_30d, pd_60d, pd_90d, risk_label, stage_migration_prob, dpd_escalation_prob, recommended_action', '15 minutes'],
        ['shap_explanations.csv', 'top_factor_1/2/3, shap_1/2/3 (SHAP values for top 3 risk drivers)', '15 minutes'],
        ['risk_flags.csv', '6 binary flags: zscore_anomaly, score_deterioration, exposure_spike, salary_stopped, overdraft_dependent, card_acceleration', '15 minutes'],
    ],
    col_widths=[1.8, 3.0, 1.4]
)

add_heading(doc, '7.2  Model Features Analyzed', 3)
features_ml = [
    ('exposure_growth_rate', 'How fast the outstanding balance is growing'),
    ('stage_age_months', 'How long in the current IFRS stage'),
    ('rating_deterioration', 'Pace of credit quality decline'),
    ('dpd_trend', 'Trajectory of days past due over time'),
    ('cure_rate', 'How often client self-cures after falling overdue'),
    ('dti_ratio', 'Debt-to-income ratio'),
    ('missed_payments', 'Count of missed payment events'),
    ('ltv_ratio', 'Loan-to-value (collateral coverage)'),
    ('overdraft_utilisation', 'Overdraft facility usage intensity'),
    ('card_utilisation', 'Credit card usage relative to limit'),
    ('salary_inflow_drop', 'Reduction in salary credits to account'),
    ('consecutive_lates', 'Streak of consecutive late payment months'),
]
for factor, desc in features_ml:
    add_bullet(doc, f' — {desc}', bold_prefix=factor)

add_heading(doc, '7.3  Composite Real-Time Risk Score (0–100)', 3)
add_para(doc, 'Even between nightly ML runs, SPECTRA computes a live composite score:')
add_code(doc, 'Score = (PD × 60) + (min(DPD, 90)/90 × 25) + (min(EWI_flags, 6)/6 × 15)')
add_bullet(doc, '60% weight on ML PD score — primary driver')
add_bullet(doc, '25% weight on DPD — operational backstop')
add_bullet(doc, '15% weight on EWI breadth — how many behavioural flags are active')
add_bullet(doc, 'Score labels: Critical ≥ 85, High ≥ 65, Medium ≥ 40, Low < 40')

add_heading(doc, '7.4  IFRS 9 Stage Classification Logic', 3)
add_table(doc,
    ['Stage', 'Condition', 'Regulatory Basis'],
    [
        ['Stage 3 (NPL)', 'DPD ≥ 90 days', 'IFRS 9 §5.5.3 — 90-day NPL definition'],
        ['Stage 3 (NPL)', 'Mortgage: DPD ≥ 60 AND PD ≥ 66%', 'Secured collateral at risk — lower threshold'],
        ['Stage 2 (SICR)', 'PD ≥ 20%', 'IFRS 9 B5.5.1 — quantitative SICR trigger'],
        ['Stage 2 (SICR)', 'DPD ≥ 30 days', 'IFRS 9 B5.5.19 — rebuttable presumption'],
        ['Stage 2 (SICR)', 'Missed payments ≥ 2', 'Qualitative SICR backstop'],
        ['Stage 2 (SICR)', 'Salary stopped + chronic overdraft', 'Combined qualitative signal'],
        ['Stage 2 (SICR)', 'Stage migration probability ≥ 40%', 'Model-based forward-looking trigger'],
        ['Stage 1', 'None of the above', 'Performing — 12-month ECL only'],
    ],
    col_widths=[1.4, 2.4, 2.4]
)
add_callout(doc,
    'Stage never auto-downgrades. Once a client reaches Stage 2 or 3, the stage can only be '
    'reduced by a manual credit committee decision — not by the system.',
    bg='FEF3C7', label='IFRS 9 Principle:')

add_heading(doc, '7.5  Key Thresholds (config.ts)', 3)
add_table(doc,
    ['Threshold', 'Default', 'Description'],
    [
        ['SICR.PD_THRESHOLD', '0.20 (20%)', 'PD level triggering Stage 1 → 2'],
        ['SICR.DPD_BACKSTOP', '30 days', 'DPD level triggering Stage 1 → 2'],
        ['SICR.NPL_DPD', '90 days', 'DPD level triggering Stage 2 → 3'],
        ['SICR.MORTGAGE_DPD', '60 days', 'Mortgage DPD level for Stage 3'],
        ['TIER.CRITICAL_PD', '0.66 (66%)', '"Default Imminent" tier threshold'],
        ['STRESS.LGD', '0.45 (45%)', 'Basel III unsecured retail Loss Given Default'],
        ['STRESS.ADVERSE_MULTIPLIER', '1.5×', 'Adverse scenario PD shock'],
        ['STRESS.SEVERE_MULTIPLIER', '2.5×', 'Severe scenario PD shock'],
        ['KPI.NPL_RED', '5%', 'NPL ratio red alert threshold'],
        ['CONCENTRATION.HHI_CONCENTRATED', '1,500', 'HHI moderate concentration warning'],
        ['CONCENTRATION.HHI_HIGHLY_CONCENTRATED', '2,500', 'HHI regulatory concern level'],
    ],
    col_widths=[2.2, 1.4, 2.6]
)

# ── 8. API Endpoints ─────────────────────────────────────────────────────────
add_heading(doc, '8.  API Endpoints Reference', 2)

add_table(doc,
    ['Method', 'Endpoint', 'Purpose', 'Role Required'],
    [
        ['POST', '/api/auth/login', 'Authenticate user → JWT cookie', 'Public'],
        ['POST', '/api/auth/logout', 'Clear session cookie', 'Any'],
        ['GET', '/api/clients/search?q=', 'Search clients or list top at-risk', 'Any'],
        ['POST', '/api/clients/[id]/insights', 'Generate AI risk analysis via Claude', 'Any'],
        ['POST', '/api/clients/[id]/action', 'Log RM action (Freeze, Escalate, etc.)', 'Any'],
        ['POST', '/api/ewi/fire', 'Fire EWI signal → trigger reclassification', 'Any'],
        ['GET/PATCH', '/api/monitoring/[id]', 'Get / update monitoring state and freeze', 'Any / RO+'],
        ['GET/POST', '/api/monitoring/[id]/documents', 'List / create document requests', 'Any'],
        ['PATCH', '/api/monitoring/[id]/documents/[docId]', 'Mark document Received', 'Any'],
        ['GET/POST', '/api/monitoring/[id]/collateral', 'List / create collateral reviews', 'Any'],
        ['GET/POST', '/api/clients/[id]/engagements', 'List / create call or meeting records', 'Any'],
        ['PATCH', '/api/clients/[id]/engagements/[engId]', 'Update engagement status/outcome', 'Any'],
        ['GET/POST', '/api/clients/[id]/covenant-waivers', 'List / request covenant waivers', 'Any'],
        ['PATCH', '/api/clients/[id]/covenant-waivers/[id]', 'Approve or reject a waiver', 'RO+ only'],
        ['GET/POST', '/api/clients/[id]/restructuring', 'List / propose restructuring plan', 'Any'],
        ['PATCH', '/api/clients/[id]/restructuring/[id]', 'Approve / advance / reject plan', 'RO+ for Approve'],
        ['GET', '/api/notifications', 'Fetch RM notification inbox', 'Any'],
        ['PUT', '/api/notifications/[id]/read', 'Mark notification as read', 'Any'],
        ['GET', '/api/db/ping', 'Database connection health check', 'Admin'],
        ['POST', '/api/cache/invalidate', 'Flush in-process result cache', 'Admin'],
    ],
    col_widths=[0.7, 2.5, 2.0, 1.0]
)
add_para(doc, 'RO+ = risk_officer or admin role required', italic=True, size=9, color=(100, 116, 139))

# ── 9. Use Cases ─────────────────────────────────────────────────────────────
add_heading(doc, '9.  Use Cases', 2)

use_cases = [
    ('Retail Bank — Portfolio Monitoring',
     'An RM opens the dashboard each morning and sees 3 new Stage 2 clients flagged overnight. '
     'The notification bell shows 5 unread alerts. The RM clicks through, reviews SHAP-explained PD scores, '
     'logs calls, and proposes restructuring plans — all within SPECTRA.'),
    ('Microfinance — Document Collection',
     'A credit officer needs financial health verification for 12 clients ahead of a credit committee. '
     'They create Document Requests for financial statements and tax returns, track Pending → Received '
     'status per client, and present outstanding ones at the committee.'),
    ('Credit Risk Team — Stress Testing',
     'The risk team presents ECL impact under three scenarios before a regulatory submission. '
     'SPECTRA applies configurable PD multipliers (Adverse: 1.5×, Severe: 2.5×) and shows projected ECL, '
     'provision coverage, and at-risk exposure by segment.'),
    ('Fintech Lender — Automated EWI',
     'The bank\'s transaction monitoring fires salary_stopped via POST /api/ewi/fire. '
     'SPECTRA immediately re-evaluates the client\'s stage, writes an audit record, and notifies '
     'the assigned RM — all without any manual intervention.'),
    ('Risk Officer — Covenant Waiver Approval',
     'An RM raises a financial_covenant waiver request on a client profile. The risk officer '
     'receives a notification, reviews the request, and approves it with timestamped decision notes.'),
]

for title_uc, desc_uc in use_cases:
    add_bullet(doc, f' — {desc_uc}', bold_prefix=title_uc)

# ── 10. Outputs & Insights ───────────────────────────────────────────────────
add_heading(doc, '10.  Outputs & Insights', 2)

add_heading(doc, 'Per-Client Profile', 3)
add_table(doc,
    ['Tab / Section', 'What It Shows'],
    [
        ['Header', 'Name, stage badge, region, employment, total exposure, ML PD (90d)'],
        ['KPI Strip', 'Credit utilisation, DPD, risk score (0–10), DTI ratio, active actions count'],
        ['Overview — Loan Products', 'All credit accounts with Stage, DPD, and approved amount'],
        ['Overview — Restructuring Plan', 'Active plan type, status, terms, and proposed date'],
        ['Overview — DPD History Chart', 'Month-by-month DPD bar chart (last 12 months)'],
        ['Overview — ML Prediction', 'PD bars at 30/60/90d, stage migration prob, SHAP top 3 drivers'],
        ['EWI Signals', 'Salary, overdraft, card, consecutive lates — severity-coded'],
        ['Alerts', 'Per-product overdue alerts with priority and recommended action'],
        ['AI Insights', 'Risk narrative, deterioration prediction, recovery strategy, transparency letter'],
        ['Actions Log', 'Full timeline of system and manual actions with timestamps'],
        ['Sidebar', 'Quick action buttons, Stage 2/3 transparency letter callout'],
    ],
    col_widths=[2.2, 4.0]
)

add_heading(doc, 'Portfolio-Level Outputs', 3)
add_bullet(doc, 'Dashboard: NPL ratio, Stage 2 rate, delinquency rate, Health Score, IFRS donut chart')
add_bullet(doc, 'Portfolio: Exposure by product and region, top 8 loans by exposure')
add_bullet(doc, 'Analytics: Rollrate matrix (5×5 DPD transition), ECL provision gap, vintage analysis, interest at risk')
add_bullet(doc, 'Concentration: HHI by product and region, Lorenz curve, top 15 obligors')
add_bullet(doc, 'Stress Test: Three scenario comparison (Base / Adverse / Severe), PD migration chart')
add_bullet(doc, 'Model: AUC-ROC, model age (Fresh/Recent/Stale), top 25 clients by PD score')
add_bullet(doc, 'Audit: Action log with user breakdown, top 6 action types, active freeze count')

# ── 11. Benefits ─────────────────────────────────────────────────────────────
add_heading(doc, '11.  Benefits', 2)

add_heading(doc, 'For Risk Management', 3)
add_bullet(doc, 'IFRS 9 compliance built in — stages updated continuously, not in month-end batch')
add_bullet(doc, 'No staging surprises at reporting date — classification runs daily')
add_bullet(doc, 'Audit-ready — every classification, RM action, and system event logged immutably')
add_bullet(doc, 'Earlier intervention — catching Stage 1 → 2 migration before it becomes Stage 3 reduces LGD')

add_heading(doc, 'For Relationship Managers', 3)
add_bullet(doc, 'Single workspace — no switching between spreadsheets, email, and CRM')
add_bullet(doc, 'Prioritized workload — system tells RM which clients need urgent attention and what to do')
add_bullet(doc, 'Engagement history — every call and meeting logged and searchable')

add_heading(doc, 'For Credit Committees', 3)
add_bullet(doc, 'Quantitative backing — ML PD scores and SHAP explanations replace gut-feel-only decisions')
add_bullet(doc, 'Restructuring proposals with full terms and approval lifecycle tracking')
add_bullet(doc, 'Stress test outputs ready for regulatory reporting')

add_heading(doc, 'For the Institution', 3)
add_bullet(doc, 'Scalable — handles large portfolios without adding headcount')
add_bullet(doc, 'Configurable — all thresholds are environment-variable-driven, no code changes needed')
add_bullet(doc, 'Reduced provisioning surprises — ECL estimates more accurate with continuous monitoring')

# ── 12. Example Scenario ─────────────────────────────────────────────────────
add_heading(doc, '12.  Example Scenario — Detecting a Risky Client', 2)

add_para(doc,
    'Client Ahmed Berisha (ID: AHM-001932) has a mortgage and a personal loan. '
    'He was Stage 1 with PD 12% three months ago.')

doc.add_paragraph()
add_heading(doc, 'Week 1 — ML Pipeline Detects Deterioration', 3)
add_table(doc,
    ['Signal', 'Value', 'Meaning'],
    [
        ['pd_90d', '0.38 (up from 0.12)', 'PD tripled — SICR quantitative trigger crossed'],
        ['flag_salary_stopped', 'True', 'No salary credit in 45 days'],
        ['flag_overdraft_dependent', 'True', 'Account in overdraft 38 of last 45 days'],
        ['stage_migration_prob', '0.44', 'Model predicts 44% probability of stage migration'],
        ['top SHAP factor', 'salary_inflow_drop (+0.187)', 'Primary driver of PD increase'],
    ],
    col_widths=[2.0, 1.8, 2.4]
)

add_heading(doc, 'Classification Engine Response', 3)
add_bullet(doc, 'PD 0.38 ≥ 0.20 SICR threshold ✓')
add_bullet(doc, 'Stage migration probability 0.44 ≥ 0.40 threshold ✓')
add_bullet(doc, 'Salary stopped + chronic overdraft ✓')
add_bullet(doc, 'Decision: Stage 1 → Stage 2', bold_prefix='')
add_code(doc, '→ recordSystemAction({ old_stage: 1, new_stage: 2, trigger_reason: {...} })')
add_code(doc, '→ createNotification({ priority: "high", assigned_rm: "maria.kovac" })')
add_code(doc, '→ upsertClientMonitoring({ review_frequency: "Weekly", is_freezed: 1 })')

add_heading(doc, 'Week 2 — RM Takes Action', 3)
actions_taken = [
    'Schedules a call → ClientEngagement record created (type: call, status: scheduled)',
    'Requests financial statement → DocumentRequest created (status: Pending)',
    'After call (outcome: reached) → updates engagement to completed',
    'Generates AI Insights → receives risk narrative + transparency letter',
    'Proposes Payment Holiday restructuring plan → 3 months, linked to mortgage',
]
for a in actions_taken:
    add_bullet(doc, a)

add_heading(doc, 'Week 6 — Resolution', 3)
add_para(doc,
    'Ahmed resumes salary payments. PD drops to 0.14. Stage remains 2 (no auto-downgrade). '
    'Risk officer approves the restructuring plan (status: Active). The plan is visible in '
    'Ahmed\'s Overview tab. Audit trail covers: 1 SystemAction, 2 ClientEngagements, '
    '1 DocumentRequest, 1 RestructuringPlan, multiple ClientActions — all attributed and timestamped.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PART II — PAGE-BY-PAGE COMPUTATION REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'PART II — PAGE-BY-PAGE COMPUTATION REFERENCE', 1)
add_divider(doc)

# ── 13. Foundation Layer ─────────────────────────────────────────────────────
add_heading(doc, '13.  Foundation Layer — DB, Cache, and Config', 2)

add_heading(doc, '13.1  Database Connection (db.ts)', 3)
add_para(doc,
    'Every piece of data flows through query<T>(sqlText, params, timeoutMs). '
    'The function: checks if the pool exists and is connected; if not, creates it using either '
    'Windows auth (msnodesqlv8 + ODBC connection string) or SQL auth (mssql/tedious); '
    'binds all parameters via req.input() — never string-concatenation; '
    'executes the query with a per-request timeout; '
    'on error, logs the first 3 SQL lines and re-throws.')

add_heading(doc, '13.2  Caching System (queries.ts)', 3)
add_bullet(doc, 'Dedicated date caches: _maxCalcDate (latest RiskPortfolio date), _maxDateID (latest DueDaysDaily date), _prevCalcDate (second-most-recent date for MoM comparisons)')
add_bullet(doc, 'In-flight coalescing: if two concurrent requests hit a cold cache, both await the same Promise — only one DB query fires')
add_bullet(doc, 'Generic LRU result cache: 200-entry max, evicts oldest on overflow, 5-minute TTL for most queries, 10-minute TTL for EWI aggregates')

add_heading(doc, '13.3  Configuration Thresholds (config.ts)', 3)
add_para(doc,
    'All business-rule thresholds are centralized in config.ts and overridable via environment variables '
    'without code changes. See Section 7.5 for the full threshold table. IFRS 9 / Basel III parameters '
    'require risk committee sign-off before modification.')

# ── 14. Dashboard Page ───────────────────────────────────────────────────────
add_heading(doc, '14.  Dashboard Page ( / )', 2)

add_heading(doc, 'getDashboardKPIs() — KPI Card Numbers', 3)
add_bullet(doc, 'Total Clients: COUNT(DISTINCT PersonalID) from DueDaysDaily on the latest snapshot date')
add_bullet(doc, 'Delinquency Rate: (clients with DPD ≥ 30 / total clients) × 100 — red ≥ 10%')
add_bullet(doc, 'Avg DPD: AVG(DueDays) across all clients on latest snapshot')
add_bullet(doc, 'Total Exposure: SUM(totalExposure) from RiskPortfolio — displayed as €X.XM')
add_bullet(doc, 'NPL Ratio: (Stage 3 exposure / total exposure) × 100 — red ≥ 5%, amber ≥ 3%')
add_bullet(doc, 'Health Score: SQL computed as 100 − (delinquency × 2) − (npl × 5). Label: Healthy ≥ 80, Moderate ≥ 60, At Risk < 60')

add_heading(doc, 'IFRS 9 Stage Donut Chart', 3)
add_code(doc, 'total = stages.reduce((s, x) => s + x.count, 0)')
add_code(doc, 'circ  = 2 × π × 46         // circumference of SVG circle (radius 46)')
add_code(doc, 'dash  = (count / total) × circ   // filled arc length per stage')
add_code(doc, 'pct   = (count / total) × 100    // percentage label')

add_heading(doc, 'Monthly Exposure Trend', 3)
add_para(doc,
    'SQL groups by LEFT(CalculationDate, 7) — extracts "YYYY-MM" — and sums totalExposure per month '
    'over the past 12 months. The line chart connects each month\'s aggregate exposure. '
    'Wrapped in its own try/catch — fails gracefully if no historical data exists.')

# ── 15. Portfolio Page ───────────────────────────────────────────────────────
add_heading(doc, '15.  Portfolio Page ( /portfolio )', 2)

add_heading(doc, 'Stage Percentage Bar', 3)
add_para(doc,
    'Three stage percentages computed in SQL: stage1_pct = (Stage 1 exposure / total) × 100, '
    'same for Stage 2 and 3. The stacked bar uses CSS flex with flex: pct — '
    'so Stage 1 at 75% gets flex: 75, Stage 2 at 20% gets flex: 20, etc.')

add_heading(doc, 'Exposure by Region — Delinquency Bar', 3)
add_code(doc, 'bar_width = Math.min(delinquency_pct × 5, 100) + "%"')
add_para(doc,
    'Scaled ×5 for visual clarity (20% delinquency fills the bar, not 100%). '
    'Capped at 100% to prevent overflow. Color: red ≥ 8%, amber ≥ 6%, green otherwise.')

add_heading(doc, 'Exposure by Product', 3)
add_para(doc,
    'SQL uses CROSS JOIN to a totals CTE: this makes the grand total available on every row '
    'without a subquery per row. pct = (segment_exposure / grand_total) × 100.')

# ── 16. Clients List Page ────────────────────────────────────────────────────
add_heading(doc, '16.  Clients List Page ( /clients )', 2)

add_heading(doc, 'Search Flow', 3)
add_bullet(doc, '280ms debounce: user stops typing → delayed fetch fires')
add_bullet(doc, 'Query ≥ 2 chars: searchClients(q) → LIKE search on PersonalID and full name')
add_bullet(doc, 'No query: getHighRiskClientsList() → top 100 by Stage DESC, totalExposure DESC')
add_bullet(doc, 'EWI filter: getEWIFilteredClients(ewi) → 4 variants (salary_stopped, overdraft, card_high, consec_lates)')
add_bullet(doc, 'Parallel: getFrozenClientIds() always runs alongside the search — merges frozen: true onto matching rows')

add_heading(doc, 'Frontend Derivations', 3)
add_code(doc, 'DPD bar width = Math.min((current_due_days / 90) × 100, 100) + "%"')
add_code(doc, 'Initials = full_name.split(" ").map(w => w[0]).slice(0, 2).join("")')
add_code(doc, 'Stage filter = results.filter(c => c.stage === selectedStage)')

# ── 17. Client Profile Page ──────────────────────────────────────────────────
add_heading(doc, '17.  Client Profile Page ( /clients/[id] )', 2)

add_heading(doc, 'getClientProfile() — The Master Query', 3)
add_para(doc, 'Ten CTEs in a single query compute:')
add_table(doc,
    ['Field', 'Computation'],
    [
        ['age', 'DATEDIFF(YEAR, DOB, GETDATE())'],
        ['on_balance', 'totalExposure × 0.85 (estimated 85% drawn)'],
        ['off_balance', 'totalExposure × 0.15 (estimated 15% undrawn)'],
        ['exposure_growth_pct', '(current_exposure − prev_exposure) / prev_exposure × 100'],
        ['dti_ratio', '(totalExposure × 0.03 / avg_monthly_income) × 100'],
        ['repayment_rate_pct', '(total_payments − missed_payments) / total_payments × 100'],
        ['tenure_years', 'DATEDIFF(YEAR, first_credit_date, GETDATE())'],
        ['risk_score (0–10)', 'Stage_base (1/4/7) + DPD_bucket (0/1/2) + missed_ratio (0–1), capped at 10'],
        ['risk_tier', 'DPD ≥ 90 or Stage 3 → default-imminent; DPD ≥ 30 or Stage 2 → deteriorating; else stable-watch'],
        ['sicr_flagged', 'Stage ≥ 2 OR DPD ≥ 30'],
    ],
    col_widths=[2.2, 4.0]
)

add_heading(doc, 'KPI Strip (Frontend)', 3)
add_code(doc, 'creditUtil = Math.round((on_balance / approved_amount) × 100)')
add_code(doc, '// Red > 80%, amber > 60%, green otherwise')
add_code(doc, 'alertCount = activeActions.length')
add_code(doc, 'derivedAlerts = products.filter(p => p.due_days > 0)')

add_heading(doc, 'ML Badge', 3)
add_code(doc, 'pd = prediction?.pd_90d  →  display as (pd × 100).toFixed(1) + "%"')
add_code(doc, 'pdColor: red ≥ 0.66, amber ≥ 0.21, green otherwise')

add_heading(doc, 'SHAP Drivers', 3)
add_code(doc, 'max = Math.abs(shap.shap_1)       // largest factor = 100% bar width')
add_code(doc, 'pct = Math.abs(d.shap / max) × 100  // relative bar width')
add_code(doc, 'color = shap ≥ 0 ? red (risk-increasing) : green (risk-reducing)')

add_heading(doc, 'AI Insights Generation', 3)
add_bullet(doc, 'POST /api/clients/[id]/insights — sends profile + prediction data to Claude API')
add_bullet(doc, 'Model: claude-haiku-4-5-20251001, max_tokens: 1800')
add_bullet(doc, 'Response parsed with regex /\\{[\\s\\S]*\\}/ to extract JSON from model output')
add_bullet(doc, 'Returns: risk_narrative, deterioration_prediction, recommended_actions, recovery_recommendation, transparency_letter')

# ── 18. Analytics Page ───────────────────────────────────────────────────────
add_heading(doc, '18.  Analytics Page ( /analytics )', 2)

add_heading(doc, 'ECL Calculation per Stage', 3)
add_table(doc,
    ['Stage', 'PD Used', 'LGD', 'ECL Formula'],
    [
        ['Stage 1', '1% (12-month)', '45%', 'exposure × 0.01 × 0.45'],
        ['Stage 2', '20% (lifetime)', '45%', 'exposure × 0.20 × 0.45'],
        ['Stage 3', '100% (credit-impaired)', '45%', 'exposure × 1.00 × 0.45'],
    ],
    col_widths=[1.2, 1.8, 1.0, 2.2]
)
add_para(doc, 'provision_gap = calculated_ecl − bank_provision. Negative = under-provisioned (red alert).',
         italic=True, size=10, color=(71, 85, 105))

add_heading(doc, 'Rollrate Matrix (5×5 DPD Bucket Transitions)', 3)
add_para(doc,
    'SQL joins DueDaysDaily on two consecutive snapshot dates. For each client, it records '
    'which DPD bucket they were in (Current / 1-29 / 30-59 / 60-89 / 90+) on each date, '
    'counts transitions, and divides by row totals to produce rate_pct. '
    'Frontend: intensity = min(pct / 30, 1) — 30% fills the cell fully. '
    'Red for moves to worse buckets, green for improvement.')

add_heading(doc, 'Repayment Summary', 3)
add_code(doc, 'repayment_ratio = SUM(OTPLATA) / SUM(ANUITET)    // paid / scheduled')
add_code(doc, 'full_pct     = ratio ≥ 1.0   → fully paid')
add_code(doc, 'partial_pct  = 0.5 ≤ ratio < 1.0  → partial payment')
add_code(doc, 'critical_pct = ratio < 0.5   → paying less than half — red alert if > 20%')

add_heading(doc, 'Interest at Risk', 3)
add_code(doc, 'interest_income_at_risk = SUM(totalExposure) × AVG(Interes) / 100')
add_para(doc, 'Computed for Stage 2 and 3 only — the annual interest income the bank may not collect.')

# ── 19. Watchlist Page ───────────────────────────────────────────────────────
add_heading(doc, '19.  Watchlist Page ( /watchlist )', 2)

add_heading(doc, 'Data Source', 3)
add_para(doc,
    'All clients where ClientActions.action IN (\'Add to Watchlist\', \'Add to watchlist\') '
    'AND status = \'active\'. Live risk data (stage, exposure, DPD) joined from RiskPortfolio and '
    'DueDaysDaily. days_on_watch = DATEDIFF(DAY, createdAt, GETDATE()).')

add_heading(doc, 'Frontend Aggregations', 3)
add_code(doc, 'totalExposure = clients.reduce((s, c) => s + c.exposure, 0)')
add_code(doc, 'avgDPD = Math.round(sum(current_due_days) / clients.length)')
add_code(doc, 'overdue   = clients.filter(c => c.days_on_watch > 60).length    // red')
add_code(doc, 'reviewDue = clients.filter(c => c.days_on_watch 30-60).length   // amber')
add_code(doc, 'current   = clients.filter(c => c.days_on_watch <= 30).length   // green')

# ── 20. Stress Test Page ─────────────────────────────────────────────────────
add_heading(doc, '20.  Stress Test Page ( /stress )', 2)

add_para(doc,
    'No SQL queries — pure JavaScript computation on predictions.csv data (15-min cached).')

add_heading(doc, 'Scenario Computation', 3)
add_code(doc, 'shocked_pd = Math.min(pd_base × multiplier, 1.0)   // cap at 100%')
add_code(doc, 'avgPD = sum(shocked_pd) / total_clients')
add_code(doc, 'elr   = avgPD × STRESS.LGD (0.45)                  // Expected Loss Rate')
add_code(doc, 'criticalCount = shocked_pd.filter(pd => pd >= 0.66).length')
add_table(doc,
    ['Scenario', 'PD Multiplier', 'Meaning'],
    [
        ['Base', '1.0×', 'Current portfolio, no shock'],
        ['Adverse', '1.5× (default)', 'Moderate downturn — GDP approximately −2%'],
        ['Severe', '2.5× (default)', 'Systemic crisis — GDP approximately −5% or worse'],
    ],
    col_widths=[1.5, 1.5, 3.2]
)

# ── 21. Concentration Risk Page ──────────────────────────────────────────────
add_heading(doc, '21.  Concentration Risk Page ( /concentration )', 2)

add_heading(doc, 'Herfindahl-Hirschman Index (HHI)', 3)
add_code(doc, 'HHI = Σ (segment_share_pct²)')
add_code(doc, '// computed in SQL: SUM(POWER(100.0 × segment_exposure / total, 2))')
add_code(doc, 'HHI < 1,500  → Dispersed (safe)')
add_code(doc, 'HHI 1,500-2,500 → Moderately Concentrated (monitor)')
add_code(doc, 'HHI > 2,500  → Highly Concentrated (regulatory concern — EBA/GL/2018/06)')
add_para(doc,
    'Window function SUM(SUM(...)) OVER () computes the grand total in one pass inside '
    'the CTE — no subquery needed per row.')

add_heading(doc, 'Top Obligors', 3)
add_code(doc, 'top1Pct    = obligors[0].pct_of_portfolio')
add_code(doc, 'top10Pct   = obligors.slice(0, 10).reduce((s, o) => s + o.pct, 0)  // rounded to 1 dp')
add_code(doc, 'largeCount = obligors.filter(o => o.pct >= 2.0).length  // 2% = large exposure floor')

add_heading(doc, 'Lorenz Curve', 3)
add_para(doc,
    'Sort obligors by exposure descending. Accumulate exposure cumulatively. '
    'Each point: x = (rank / total_clients × 100), y = (cumulative_exposure / total × 100). '
    'SVG polygon filled between the curve and the 45° diagonal. '
    'The gap represents portfolio inequality — a straight 45° line = perfect equality.')

# ── 22. Model Intelligence Page ──────────────────────────────────────────────
add_heading(doc, '22.  Model Intelligence Page ( /model )', 2)

add_heading(doc, 'Model Age', 3)
add_code(doc, 'modelAge = Math.round((Date.now() - new Date(training_date)) / 86_400_000)')
add_code(doc, 'label: "Fresh" ≤ 7d | "Recent" ≤ 30d | "Stale" > 30d')

add_heading(doc, 'AUC-ROC Display', 3)
add_code(doc, 'auc from training_meta.json → targets.label_default_90d.auc')
add_code(doc, 'bar_width = auc × 100 + "%"')
add_code(doc, 'label: Excellent ≥ 0.85 | Good ≥ 0.75 | Fair ≥ 0.65 | Weak < 0.65')
add_code(doc, 'color: green ≥ 0.85 | amber ≥ 0.75 | red otherwise')

add_heading(doc, 'Top Risk Clients Table', 3)
add_code(doc, 'predictions.sort((a, b) => b.pd_score - a.pd_score).slice(0, 25)')
add_code(doc, 'pd_bar_width = (pd_score × 100).toFixed(1) + "%"')

# ── 23. Audit Log Page ───────────────────────────────────────────────────────
add_heading(doc, '23.  Audit Log Page ( /audit )', 2)

add_heading(doc, 'getAuditStats()', 3)
add_code(doc, 'total_today  = COUNT(*) WHERE createdAt date = today')
add_code(doc, 'total_week   = COUNT(*) WHERE createdAt >= 7 days ago')
add_code(doc, 'active_freezes = COUNT(*) WHERE action IN ("Freeze Account") AND status = "active"')

add_heading(doc, 'Frontend Breakdowns', 3)
add_code(doc, '// Action type breakdown')
add_code(doc, 'actionCounts[action]++  for each log entry')
add_code(doc, 'topActions = Object.entries(actionCounts).sort by count DESC.slice(0, 6)')
add_code(doc, '// User breakdown')
add_code(doc, 'byUser[actionedBy]++  for each log entry')
add_code(doc, 'topUsers = Object.entries(byUser).sort by count DESC.slice(0, 5)')

# ── 24. Classification Engine ────────────────────────────────────────────────
add_heading(doc, '24.  The Classification Engine (classificationEngine.ts)', 2)

add_heading(doc, 'Composite Risk Score', 3)
add_code(doc, 'pdComponent  = pdScore × 60          // 60 pts max')
add_code(doc, 'dpdComponent = (min(DPD, 90) / 90) × 25   // 25 pts max')
add_code(doc, 'ewiComponent = (min(ewiFlagCount, 6) / 6) × 15   // 15 pts max')
add_code(doc, 'score = min(100, round(pdComponent + dpdComponent + ewiComponent))')

add_heading(doc, 'Stage Derivation Order', 3)
add_bullet(doc, 'Stage 3 check first: currentStage ≥ 2 AND DPD ≥ 90 → Stage 3')
add_bullet(doc, 'Mortgage Stage 3: currentStage ≥ 2 AND isMortgage AND PD ≥ 0.66 AND DPD ≥ 60 → Stage 3')
add_bullet(doc, 'SICR triggers (any one → Stage 2): PD ≥ 0.20, DPD ≥ 30, missed ≥ 2, salary+overdraft, stageMigProb ≥ 0.40')
add_bullet(doc, 'No triggers → stay at currentStage (never auto-downgrade)')

add_heading(doc, 'Persistence Logic', 3)
add_code(doc, 'shouldPersist = stageChanged OR abs(newScore - oldScore) >= 5')
add_code(doc, 'if (shouldPersist) → recordSystemAction()')
add_code(doc, 'if (stageChanged) → createNotification() + upsertClientMonitoring()')

# ── 25. Action Engine ────────────────────────────────────────────────────────
add_heading(doc, '25.  The Action Engine (actionEngine.ts)', 2)

add_table(doc,
    ['Urgency', 'Trigger', 'Action'],
    [
        ['IMMEDIATE', 'Stage 3 or DPD ≥ 90', 'Escalate → Recovery + Legal Review'],
        ['IMMEDIATE', 'Mortgage DPD ≥ 60', 'Legal Review'],
        ['IMMEDIATE', 'PD ≥ 0.70 + DPD ≥ 30 + salary stopped', 'Freeze Account'],
        ['URGENT', 'PD ≥ 0.50 + 30 ≤ DPD < 90', 'Call Now'],
        ['URGENT', 'Salary stopped', 'Request salary documentation'],
        ['URGENT', 'stageMigProb ≥ 0.40', 'Add to Watchlist'],
        ['URGENT', 'PD ≥ 0.50 + DPD=0 + escalProb ≥ 0.40', 'Schedule Call'],
        ['STANDARD', 'Stage 2 + (DPD ≥ 15 or escalProb ≥ 0.25)', 'Restructure'],
        ['STANDARD', 'Chronic overdraft', 'Review overdraft facility'],
        ['STANDARD', 'DTI ≥ 55%', 'Debt restructuring consultation'],
        ['ROUTINE', '0.25 ≤ PD < 0.50', 'Add to Watchlist'],
        ['ROUTINE', 'Stage 2', 'Flag for Review'],
        ['ROUTINE', '0.15 ≤ PD < 0.25', 'Monthly Monitor'],
    ],
    col_widths=[1.2, 2.8, 2.2]
)
add_para(doc,
    'Output: deduplicated, sorted by urgency, max 5 actions returned. '
    'Already-active actions (in the activeActions set) are never re-recommended.',
    italic=True, size=10, color=(71, 85, 105))

# ── 26. Notification Service ─────────────────────────────────────────────────
add_heading(doc, '26.  The Notification Service (notificationService.ts)', 2)

add_para(doc,
    'Two DDL-on-first-use tables: SystemActions (immutable automated audit log) and '
    'Notifications (RM inbox). Notifications are filtered to the requesting user: '
    'WHERE assigned_rm = @username OR assigned_rm IS NULL. '
    'NULL = broadcast to all RMs. markAllReadForUser() updates read_at = GETDATE() '
    'for all unread notifications for the user.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# PART III — ROLE-BASED USAGE GUIDES
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'PART III — ROLE-BASED USAGE GUIDES', 1)
add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# RISK OFFICER
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '27.  Risk Officer Guide', 2)

add_para(doc,
    'The Risk Officer is the primary decision-maker in SPECTRA. You have full access to all restricted '
    'actions — freezes, approvals, legal escalations, restructuring decisions — and you are the person '
    'the system notifies first when something escalates.')

add_heading(doc, '27.1  Daily Routine', 3)

add_heading(doc, 'Step 1 — Notification Bell', 4)
add_para(doc,
    'Every morning: click the top-right notification bell. The red badge shows unread count. '
    'Notifications are ordered by priority: Critical (Stage 3), High (Stage 2), Medium (risk score drift). '
    'Each shows: client ID, what changed, and the exact SICR trigger that caused it. '
    'Work top-to-bottom. Click each client name to jump directly to their profile.')

add_heading(doc, 'Step 2 — Dashboard KPI Review', 4)
add_table(doc,
    ['KPI', 'Watch For'],
    [
        ['NPL Ratio', 'Red ≥ 5%, amber ≥ 3% — if red, systemic portfolio problem'],
        ['Stage 2 Rate', 'Red ≥ 15%, amber ≥ 8% — early warning of brewing NPL'],
        ['Delinquency Rate', 'Red ≥ 10% — % of clients with DPD ≥ 30'],
        ['Health Score', '< 60 = "At Risk" — act before month-end reporting'],
        ['IFRS Donut', 'Stage 2 + Stage 3 slices = your provisioning exposure'],
    ],
    col_widths=[1.8, 4.4]
)

add_heading(doc, 'Step 3 — Work Stage 3 Clients First', 4)
add_para(doc, 'Go to /clients → filter Stage 3. For each client:')
add_numbered(doc, 'Open client profile → Alerts tab → read recommended action (will say IMMEDIATE)')
add_numbered(doc, 'Sidebar → Escalate (logs "Escalate → Recovery" in ClientActions)')
add_numbered(doc, 'If account not yet frozen → Sidebar → Freeze Account')
add_numbered(doc, 'If transparency letter needed → generate AI Insights → sidebar callout → Open Transparency Letter → Copy and send')
add_para(doc, 'Then repeat for Stage 2 clients. The amber sidebar callout appears automatically for Stage 2/3.')

add_heading(doc, '27.2  Restricted Actions You Own', 3)

add_heading(doc, 'Freezing an Account', 4)
add_bullet(doc, 'When: PD ≥ 70% + DPD ≥ 30 + salary stopped — the action engine labels this IMMEDIATE')
add_bullet(doc, 'How: Client Profile → Sidebar → "Freeze Account" → click → logs ClientAction status: active')
add_bullet(doc, 'Effect: client appears with 🔒 on Clients list; ClientMonitoring.is_freezed = 1')
add_bullet(doc, 'To unfreeze: resolve the "Freeze Account" action in the Actions Log (sets status = resolved)')
add_callout(doc,
    'The classification engine also auto-freezes clients when it escalates them to Stage 2 or 3. '
    'The manual freeze is for cases where the engine hasn\'t run yet or you want to act proactively.',
    bg='EFF6FF', label='Note:')

add_heading(doc, 'Approving a Covenant Waiver', 4)
add_bullet(doc, 'An RM (or analyst) raises the waiver request on the client profile')
add_bullet(doc, 'You receive a notification or find it via the client profile')
add_bullet(doc, 'PATCH /api/clients/[id]/covenant-waivers/[waiverId] with { status: "Approved", decisionNotes: "..." }')
add_bullet(doc, 'Only risk_officer and admin can send this PATCH — analysts get 403')
add_bullet(doc, 'Fields set: approved_by = your username, approved_at = timestamp, decision_notes = your notes')
add_bullet(doc, 'To reject: same endpoint with { status: "Rejected", decisionNotes: "reason" }')

add_heading(doc, 'Approving a Restructuring Plan', 4)
add_bullet(doc, 'Plan is created by anyone (analyst or RM) with status = Proposed')
add_bullet(doc, 'You approve: PATCH /api/clients/[id]/restructuring/[planId] with { status: "Approved" }')
add_bullet(doc, 'Only risk_officer and admin can set status = Approved')
add_bullet(doc, 'Advance to Active when plan goes live, then Completed when fulfilled')
add_bullet(doc, 'Or Reject with notes if not viable')

add_heading(doc, '27.3  Portfolio-Level Work', 3)

add_heading(doc, 'Concentration Risk (/concentration) — Weekly Check', 4)
add_bullet(doc, 'Top 1 obligor % — EBA watchlist if ≥ 10%')
add_bullet(doc, 'Top 10 combined % — watch if ≥ 50%')
add_bullet(doc, 'HHI < 1,500 = safe; 1,500–2,500 = moderate; > 2,500 = highly concentrated (report to regulator)')
add_bullet(doc, 'Lorenz curve — pronounced bow = high inequality in exposure distribution')

add_heading(doc, 'Stress Testing (/stress) — Before Credit Committee', 4)
add_bullet(doc, 'Base: current portfolio with no shock')
add_bullet(doc, 'Adverse (×1.5): moderate downturn simulation')
add_bullet(doc, 'Severe (×2.5): systemic crisis simulation')
add_bullet(doc, 'Present: Avg PD delta, ELR delta, critical client count change, risk label migration chart')

add_heading(doc, 'Analytics Deep Dive (/analytics)', 4)
add_bullet(doc, 'Rollrate Matrix: look at Current → 90+ DPD cell — clients skipping to NPL without transition is alarming')
add_bullet(doc, 'ECL Provision Gap: red = under-provisioned for that stage — must report and correct')
add_bullet(doc, 'Coverage by Stage MoM: if mom_change_pct < −2%, red "Declined" badge — cushion is thinning')
add_bullet(doc, 'Interest at Risk: annual interest income exposed for Stage 2 + 3 — income statement impact')

add_heading(doc, '27.4  Watchlist Management (/watchlist)', 3)
add_bullet(doc, '"Overdue Review" badge (> 60 days): every one needs your attention this week — escalate or resolve')
add_bullet(doc, '"Review Due" badge (30–60 days): schedule a review call this week')
add_bullet(doc, 'To remove: resolve the "Add to Watchlist" ClientAction for the client')
add_bullet(doc, 'To escalate: open client profile → generate AI Insights → present to credit committee')

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# ADMIN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '28.  Admin Guide', 2)

add_para(doc,
    'As Admin you have the same access as the Risk Officer (do everything in Section 27) plus '
    'system-level control over users, configuration, caching, and infrastructure.')

add_heading(doc, '28.1  User Management', 3)
add_para(doc,
    'Users are defined in src/lib/users.ts. Three roles exist: admin (full access), '
    'risk_officer (all restricted actions), analyst (read + basic logging only). '
    'To add or change a user: modify users.ts directly and redeploy. '
    'No UI-based user management screen exists currently.')
add_table(doc,
    ['Role', 'Capabilities'],
    [
        ['admin', 'Everything — all risk_officer capabilities plus system management'],
        ['risk_officer', 'All restricted actions: freeze, approve waivers, approve restructuring, legal review, escalation'],
        ['analyst', 'Read everything, log basic actions (call, meeting, document request, watchlist), propose plans/waivers'],
    ],
    col_widths=[1.2, 5.0]
)

add_heading(doc, '28.2  Cache Invalidation', 3)
add_para(doc, 'If you need fresh data immediately (after manual DB update or new ML run):')
add_code(doc, 'POST /api/cache/invalidate')
add_para(doc,
    'This flushes the in-process LRU Map and resets all module-level date caches. '
    'The next request re-queries the database. No UI button — use curl, Postman, or a browser extension. '
    'Normal cache TTLs: 5 minutes (portfolio/analytics), 10 minutes (EWI), 15 minutes (CSV files).')

add_heading(doc, '28.3  Database Health Check', 3)
add_code(doc, 'GET /api/db/ping')
add_para(doc, 'Returns: connection status, server name, SQL Server version. Use to diagnose "Database error" panels. '
    'If ping fails, check: DB_SERVER, DB_NAME, DB_USER, DB_PASSWORD, DB_ODBC_DRIVER environment variables in .env.local.')

add_heading(doc, '28.4  Threshold Configuration', 3)
add_para(doc,
    'All thresholds in config.ts are overridable via environment variables without code changes. '
    'Set in .env.local and restart the Node.js process.')
add_table(doc,
    ['To Change', 'Environment Variable', 'Example'],
    [
        ['SICR PD trigger', 'SICR_PD_THRESHOLD', '0.25'],
        ['NPL DPD definition', 'SICR_NPL_DPD', '60'],
        ['Adverse stress multiplier', 'STRESS_ADVERSE_MULTIPLIER', '2.0'],
        ['Cache duration', 'CACHE_TTL_MS', '300000 (5 min)'],
        ['NPL ratio red threshold', 'KPI_NPL_RED', '4'],
        ['EBA top obligor warning', 'CONC_TOP1_OBLIGOR_WARN', '8'],
        ['LGD (Basel)', 'STRESS_LGD', '0.45'],
    ],
    col_widths=[2.0, 2.4, 1.8]
)
add_callout(doc,
    'IFRS 9 and Basel III parameters (SICR_*, STRESS_LGD, CONC_HHI_*) require risk committee '
    'sign-off before changing. Document the reason in the change log.',
    bg='FEE2E2', label='⚠  Compliance Warning:')

add_heading(doc, '28.5  ML Pipeline Monitoring', 3)
add_para(doc, 'Check /model page:')
add_bullet(doc, '"Fresh" (≤ 7 days): pipeline running on schedule')
add_bullet(doc, '"Recent" (≤ 30 days): acceptable but schedule a run soon')
add_bullet(doc, '"Stale" (> 30 days): pipeline has not run — investigate Python scheduler')
add_bullet(doc, 'AUC-ROC < 0.75 (red "Weak"): model accuracy degrading — schedule a retraining')
add_para(doc, 'If CSV files are missing, SPECTRA degrades gracefully — pages still load but ML columns show "—". '
    'EWI and DPD data from SQL is unaffected.')

add_heading(doc, '28.6  SPECTRA-Owned Table Verification', 3)
add_para(doc, 'Run this in SQL Server Management Studio to verify all tables exist:')
add_code(doc,
    "SELECT name, create_date FROM sys.tables\n"
    "WHERE schema_id = SCHEMA_ID('dbo')\n"
    "  AND name IN (\n"
    "    'ClientActions','SystemActions','Notifications',\n"
    "    'ClientMonitoring','DocumentRequests','CollateralReview',\n"
    "    'ClientEngagements','CovenantWaivers','RestructuringPlans'\n"
    "  )\n"
    "ORDER BY create_date"
)
add_para(doc, 'Missing tables are created automatically on first use. No manual DDL is ever needed.')

add_heading(doc, '28.7  Audit and Compliance Oversight', 3)
add_bullet(doc, 'active_freezes: all should have legitimate reason in actions log — investigate any without clear rationale')
add_bullet(doc, 'top actions breakdown: "Freeze Account" or "Legal Review" dominating = portfolio in distress')
add_bullet(doc, 'top users breakdown: one analyst logging disproportionate actions = investigate')
add_bullet(doc, 'SystemActions table is immutable — no SPECTRA API can update or delete records — this is your compliance guarantee')

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════════
# ANALYST
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '29.  Analyst Guide', 2)

add_para(doc,
    'As analyst your role is monitoring, research, and recommendation — not decision-making. '
    'You can read everything and log basic actions, but consequential decisions '
    '(freeze, approve, legal escalation) require a risk officer or admin. '
    'Your job: find the problems, document them, surface them to the risk officer.')

add_heading(doc, '29.1  Daily Workflow', 3)

add_heading(doc, 'Step 1 — Morning Dashboard Scan', 4)
add_para(doc,
    'Open Dashboard (/). Read the five KPIs. Note any red indicators. '
    'Check if the Stage 2 or Stage 3 slice of the IFRS donut grew since yesterday — '
    'that signals overnight reclassification events.')

add_heading(doc, 'Step 2 — Client Research (/clients)', 4)
add_bullet(doc, 'No search = top 100 highest-risk clients (default view — most dangerous first)')
add_bullet(doc, 'Stage filter tabs: count badges show how many in each stage')
add_bullet(doc, 'Search ≥ 2 chars: searches PersonalID prefix AND full name substring')
add_bullet(doc, '🔒 frozen badge = risk officer has frozen this account')

add_heading(doc, 'EWI Filter Shortcuts', 4)
add_table(doc,
    ['Filter', 'Identifies', 'Use For'],
    [
        ['Salary Stopped', 'Stage ≥ 2 clients with no salary in 60 days', 'Highest priority outreach list'],
        ['Overdraft Dependent', 'Clients with negative account balance', 'Structural liquidity risk'],
        ['Card High Usage', 'Clients with card spend > threshold in 30 days', 'Overspending / distress signal'],
        ['Consecutive Lates', 'Clients with DPD > 0 on latest snapshot', 'Active delinquency list'],
    ],
    col_widths=[1.6, 2.4, 2.2]
)

add_heading(doc, 'Step 3 — Client Profile Deep Dive', 4)

add_heading(doc, 'Overview Tab', 4)
add_bullet(doc, 'KPI strip: credit utilisation (red > 80%), DPD, risk score (0–10), DTI ratio')
add_bullet(doc, 'Loan Products: check each account\'s Stage and DPD individually')
add_bullet(doc, 'Active Restructuring Plan: note type and status before making further recommendations')
add_bullet(doc, 'DPD History: rising bars = deteriorating trend; all green = performing')
add_bullet(doc, 'ML Prediction: rising PD slope (30d < 60d < 90d) = expected deterioration ahead. SHAP drivers = why.')

add_heading(doc, 'EWI Signals Tab', 4)
add_bullet(doc, 'Four indicator cards: Salary, Overdraft, Card, Consecutive Lates — red = alert')
add_bullet(doc, 'Signal Log: severity table (High / Medium / Low) — High signals are the ones to escalate')
add_bullet(doc, 'ML Risk Flags: 6 binary anomaly flags from the Python pipeline — each has a plain-English description')

add_heading(doc, 'AI Insights Tab — Your Most Powerful Tool', 4)
add_para(doc, 'Click "Generate AI Insights". Receives from Claude AI:')
add_numbered(doc, 'Risk Narrative — 2-3 sentence assessment with risk level badge and key concern')
add_numbered(doc, 'Deterioration Prediction — probability statement, risk score meter, 3 key signals')
add_numbered(doc, 'Recommended Actions — primary action + supporting list + committee escalation flag')
add_numbered(doc, 'Recovery Recommendation — probability, strategy type, estimated recovery rate')
add_numbered(doc, 'Transparency Letter — fully drafted formal letter ready to copy and send')
add_callout(doc,
    'Use AI Insights to build your case before presenting to the risk officer. '
    'The transparency letter can be copied directly — the risk officer approves and sends it.',
    bg='ECFDF5', label='Tip:')

add_heading(doc, '29.2  Actions You Can Log', 3)

add_table(doc,
    ['Action', 'Available?', 'Creates', 'Notes'],
    [
        ['Schedule Call', '✅ Yes', 'ClientEngagement record', 'Type: call, status: scheduled — has lifecycle tracking'],
        ['Schedule Meeting', '✅ Yes', 'ClientEngagement record', 'Type: meeting, status: scheduled'],
        ['Request Documents', '✅ Yes', 'ClientAction log entry', 'Use DocumentRequests API for formal tracking'],
        ['Add to Watchlist', '✅ Yes', 'ClientAction log entry', 'Client appears in /watchlist immediately'],
        ['Freeze Account', '✅ Log only', 'ClientAction log entry', 'Risk officer must confirm/enforce'],
        ['Escalate', '✅ Log only', 'ClientAction log entry', 'Actual decision made by risk officer'],
        ['Approve Waiver', '❌ No', '—', '403 Forbidden — risk_officer role required'],
        ['Approve Restructuring', '❌ No', '—', '403 Forbidden — risk_officer role required'],
    ],
    col_widths=[1.8, 1.0, 1.8, 2.6]
)

add_heading(doc, '29.3  Proposing a Restructuring Plan', 3)
add_numbered(doc, 'Client Profile → Sidebar → "Propose Restructuring"')
add_numbered(doc, 'Select Plan Type: Loan Extension / Payment Holiday / Rate Reduction / Debt Consolidation / Partial Write-Off')
add_numbered(doc, 'Fill conditional fields based on type (maturity date / months / rate / amount)')
add_numbered(doc, 'Optionally link to a specific Credit Account or leave blank for all credits')
add_numbered(doc, 'Add notes explaining the rationale')
add_numbered(doc, 'Click "Submit Proposal" — status starts as Proposed')
add_numbered(doc, 'Risk officer sees it in client\'s Overview tab and can approve or reject')

add_heading(doc, '29.4  Requesting a Covenant Waiver', 3)
add_numbered(doc, 'Identify the covenant the client cannot meet')
add_numbered(doc, 'POST /api/clients/[id]/covenant-waivers with: waiverType, requestedDate, reason, creditId (optional)')
add_numbered(doc, 'Status starts as Pending — risk officer approves or rejects')
add_para(doc, 'Waiver types: financial_covenant | reporting_covenant | maintenance_covenant | other')

add_heading(doc, '29.5  Watchlist Management', 3)
add_bullet(doc, 'Add client: Sidebar → action button or log via action engine recommendation')
add_bullet(doc, 'When: action engine recommends "Add to Watchlist" at 0.25 ≤ PD < 0.50 OR stageMigProb ≥ 40%')
add_bullet(doc, 'After 30 days: "Review Due" badge — brief the risk officer')
add_bullet(doc, 'After 60 days: "Overdue Review" badge — requires immediate decision: escalate or remove')

add_heading(doc, '29.6  Portfolio Research (Read-Only)', 3)
add_bullet(doc, '/portfolio — Where is exposure? Which regions have highest delinquency?')
add_bullet(doc, '/analytics — Vintage Analysis (which issuance year performs worst?), Segment Delinquency (worst product type?)')
add_bullet(doc, '/model — Top 25 clients by PD score = your highest-priority daily call list')
add_bullet(doc, '/concentration — Top obligors table, HHI — surface concentrations at next committee meeting')

add_heading(doc, '29.7  What You Cannot Do', 3)
add_table(doc,
    ['Action', 'Reason'],
    [
        ['Approve covenant waiver', 'PATCH endpoint returns 403 for analyst role'],
        ['Approve restructuring plan', 'Status = "Approved" requires risk_officer role (server-side check)'],
        ['Invalidate cache', 'Admin-only system function'],
        ['Override stage classification', 'Manual reclassification requires risk_officer via POST /api/classify/[id]'],
        ['View DB credentials or JWT_SECRET', 'Server-only environment variables — never sent to client'],
    ],
    col_widths=[2.2, 4.0]
)
add_para(doc,
    'If you believe a stage classification is wrong: log a "Flag for Review" action with notes, '
    'generate AI Insights to document your reasoning, and bring it to the risk officer.',
    italic=True, size=10, color=(71, 85, 105))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 30. ROLE COMPARISON TABLE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '30.  Role Comparison — Quick Reference', 2)

add_table(doc,
    ['Capability', 'Analyst', 'Risk Officer', 'Admin'],
    [
        ['View dashboard, portfolio, analytics, watchlist', '✅', '✅', '✅'],
        ['Search and view all client profiles', '✅', '✅', '✅'],
        ['Generate AI insights (risk narrative + letter)', '✅', '✅', '✅'],
        ['Schedule calls and meetings (real records)', '✅', '✅', '✅'],
        ['Propose restructuring plan', '✅', '✅', '✅'],
        ['Request covenant waiver', '✅', '✅', '✅'],
        ['Request documents', '✅', '✅', '✅'],
        ['Add to watchlist', '✅', '✅', '✅'],
        ['Log Freeze Account (intent only)', '✅', '✅', '✅'],
        ['Enforce account freeze (monitoring + classification)', '❌', '✅', '✅'],
        ['Approve / reject covenant waiver', '❌', '✅', '✅'],
        ['Approve / reject restructuring plan', '❌', '✅', '✅'],
        ['Log Legal Review', '❌', '✅', '✅'],
        ['Log Escalate → Recovery', '❌', '✅', '✅'],
        ['Trigger manual reclassification', '❌', '✅', '✅'],
        ['View and manage notifications', '✅ (broadcast)', '✅ (assigned + broadcast)', '✅ (all)'],
        ['Invalidate system cache', '❌', '❌', '✅'],
        ['Manage users and thresholds', '❌', '❌', '✅'],
        ['Database health check (/api/db/ping)', '❌', '❌', '✅'],
        ['View audit log (/audit)', '❌', '✅', '✅'],
    ],
    col_widths=[3.6, 0.9, 1.1, 0.8]
)

# ── Final footer ──────────────────────────────────────────────────────────────
doc.add_paragraph()
add_divider(doc)
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.add_run(
    f'SPECTRA Credit Risk Intelligence Platform  ·  Internal Documentation  ·  '
    f'Confidential  ·  Generated {datetime.datetime.now().strftime("%d %B %Y")}'
).font.color.rgb = RGBColor(148, 163, 184)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r'c:\Users\Elion\OneDrive\Documents\GitHub\spectra_system\SPECTRA_Complete_Documentation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
